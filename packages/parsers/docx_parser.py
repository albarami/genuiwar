"""DOCX parser — extracts paragraphs, headings, and tables from Word documents."""

from pathlib import Path

import docx

from packages.parsers.base import BaseParser, ParseError, ParseResult
from packages.parsers.registry import register_parser
from packages.schemas.document import FileDocument
from packages.schemas.enums import FileType
from packages.schemas.evidence import CitationAnchor, EvidenceChunk


@register_parser(FileType.DOCX)
class DocxParser(BaseParser):
    """Parse .docx files into evidence chunks with section-based citation anchors."""

    def parse(self, file_path: Path, file_doc: FileDocument) -> ParseResult:
        """Parse a DOCX file into evidence chunks."""
        try:
            document = docx.Document(str(file_path))
        except Exception as exc:
            raise ParseError(str(file_path), f"Cannot open DOCX: {exc}") from exc

        chunks: list[EvidenceChunk] = []
        warnings: list[str] = []
        current_section = "Document Start"

        for para in document.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                current_section = para.text.strip() or current_section
                continue

            text = para.text.strip()
            if not text:
                continue

            chunks.append(
                EvidenceChunk(
                    file_id=file_doc.file_id,
                    content=text,
                    content_type="text",
                    citation_anchor=CitationAnchor(
                        file_id=file_doc.file_id,
                        section=current_section,
                    ),
                )
            )

        for idx, table in enumerate(document.tables):
            rows_text: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))

            if rows_text:
                chunks.append(
                    EvidenceChunk(
                        file_id=file_doc.file_id,
                        content="\n".join(rows_text),
                        content_type="table",
                        citation_anchor=CitationAnchor(
                            file_id=file_doc.file_id,
                            section=current_section,
                        ),
                        metadata={"table_index": str(idx)},
                    )
                )

        if not chunks:
            warnings.append("DOCX file produced no content chunks")

        file_doc.page_count = len(document.paragraphs)

        return ParseResult(document=file_doc, chunks=chunks, warnings=warnings)
